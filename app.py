import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl.styles import PatternFill, Font
from openpyxl.utils import get_column_letter
import io
import os
import requests
import cairosvg
import traceback
import base64
from PIL import Image

# 1. Page Configuration
st.set_page_config(
    page_title="AI Readiness Tool",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for Premium Look
st.markdown("""
    <style>
    .main {
        background-color: #f8f9fa;
    }
    .stButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        background-color: #003366;
        color: white;
        font-weight: bold;
    }
    .stDownloadButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        background-color: #006699;
        color: white;
        font-weight: bold;
    }
    div[data-testid="stExpander"] {
        border: 1px solid #e9ecef;
        border-radius: 10px;
        background-color: white;
    }
    h1, h2, h3 {
        color: #003366;
    }
    </style>
    """, unsafe_allow_html=True)

# --- LOGIN MODULE ---
def check_password():
    """Returns `True` if the user had the correct password."""
    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == st.secrets.get("APP_PASSWORD", "admin123"): # Default for local dev
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # don't store password
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # First run, show input for password.
        st.text_input(
            "Podaj hasło dostępu:", 
            type="password", 
            on_change=password_entered, 
            key="password"
        )
        st.info("Dane potrzebne do zalogowania znajdują się w Monday. Kontakt: jaroslaw.muzyka@performance-group.pl")
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input(
            "Podaj hasło dostępu:", 
            type="password", 
            on_change=password_entered, 
            key="password"
        )
        st.error("😕 Niepoprawne hasło")
        st.info("Dane potrzebne do zalogowania znajdują się w Monday. Kontakt: jaroslaw.muzyka@performance-group.pl")
        return False
    else:
        # Password correct.
        return True

if not check_password():
    st.stop()

# --- DATABASE / KNOWLEDGE ---
commentary_db = {
    "Czy dodany jest plik robots.txt?": "Plik robots.txt to podstawowy mechanizm kontroli dostępu dla botów. Jego brak może prowadzić do crawlowania niepożądanych zasobów (co może wpłynąć na gorszą indeksację kluczowych podstron), a jego błędna konfiguracja może całkowicie zablokować dostęp dla botów np dla GPTBot.",
    "Czy strona jest możliwa do crawlowania przez wyszukiwarki?": "Modele AI opierają się na w dużej mierze na danych zindeksowanych przez wyszukiwarki. Jeśli strona jest niedostępna dla robotów, jej treść może być niewidoczna dla LLMów i nie zostanie wykorzystana w odpowiedziach.", 
    "Czy strona jest indeksowalna?": "Brak możliwości zindeksowania strony (np. przez tag 'noindex') może wpływać na wykluczenie jej treści z bazy wiedzy modeli AI.", 
    "Czy dodana jest mapa strony XML?": "Mapa strony to przewodnik dla robotów, ułatwiający im szybkie odnalezienie wszystkich kluczowych podstron, które powinny zostać przeanalizowane.", 
    "Czy dodawany jest znacznik <lastmod> w mapie witryny XML?": "Znacznik <lastmod> informuje o dacie ostatniej modyfikacji. Modele językowe preferują aktualne informacje, a ten znacznik może pomóc im priorytetyzować nowe treści.", 
    "Czy dodane są dedykowane mapy XML ze zdjęciami i filmami?": "Odpowiedzi AI coraz częściej zawierają multimedia. Dedykowane mapy mogą pomóc modelom AI odkrywać i poprawnie interpretować te zasoby.", 
    "Czy mapa strony XML została przesłana do Google Search Console?": "Przesłanie mapy do GSC przyspiesza proces odkrywania i indeksowania treści przez boty Google, w tym te zasilające AI.", 
    "Czy mapa strony XML została przesłana do Bing Webmaster Tools?": "Jest to kluczowe dla widoczności w ekosystemie Microsoft, w tym w wyszukiwarce Bing i Copilot.", 
    "Czy mapa strony XML jest dodana do robots.txt?": "Umieszczenie ścieżki do mapy strony w pliku robots.txt to standardowa praktyka, która ułatwia botom jej odnalezienie.", 
    "Czy mapa strony XML zawiera tylko adresy z kodem 200, kanoniczne, nie zawiera adresów noindex?": "Czysta, wolna od błędów i niepożądanych adresów mapa strony pozwala botom AI efektywniej wykorzystać czas na analizę wartościowych treści, zamiast marnować go na niepożądane przez nas adresy.", 
    "Czy plik robots.txt pozwala agentom LLM i robotom wyszukiwarek na crawlowanie strony?": "Zablokowanie botów np GPTBot może całkowicie uniemożliwiść im wykorzystanie treści na naszej stronie.", 
 
    "Czy dodany jest certyfikat SSL?": "Protokół HTTPS (SSL) jest fundamentalnym sygnałem zaufania. Strony bez certyfikatu są uznawane za niezabezpieczone co może mieć wpływ na obecność w odpowiedziach AI.", 
    "Czy strona wykorzystuje Breadcrumb?": "Breadcrumby (nawigacja okruszkowa) poza lepszym doświadczeniem użytkowników na stronie może pomagać LLM zrozumieć hierarchię i relacje między podstronami, co jest kluczowe dla generowania trafnych odpowiedzi.",
    "social_media_general": "Profile społecznościowe budują tożsamość i autorytet marki. Modele AI mogą wykorzystywać je do potwierdzenia, że marka jest prawdziwa, aktywna i jest ekspertem w swojej dziedzinie, co zwiększa jej wiarygodność. Regularna aktywność to sygnał, że informacje o marce są aktualne.",
    "linkbuilding_general": "Linki przychodzące z innych stron to jeden z najważniejszych sygnałów autorytetu w Internecie. Modele AI, podobnie jak tradycyjne wyszukiwarki, preferują informacje pochodzące z wiarygodnych, dobrze podlinkowanych źródeł.",
    "tresci_general": "Treść jest fundamentem, na którym opierają się modele AI. Musi być ona nie tylko unikalna i merytoryczna, ale także aktualna, zaufana i strukturyzowana w sposób, który ułatwia maszynom szybkie znalezienie konkretnych odpowiedzi. Elementy takie jak daty, linki do źródeł czy sekcje FAQ bezpośrednio wpływają na to, czy treść zostanie uznana za wiarygodne źródło dla generowanych odpowiedzi.",
    "non_indexable": "Modele językowe mogą nie wykorzystywać treści ze stron, które są zablokowane przed indeksowaniem (np. przez tag 'noindex'). Informacje zawarte na tych podstronach mogą być dla AI niewidoczne.", 
    "core_web_vitals": "Google traktuje Core Web Vitals jako kluczowe wskaźniki jakości i użyteczności strony. Słabe wyniki CWV mogą obniżyć postrzeganą jakość witryny i zmniejszyć jej szansę na bycie cytowaną.", 
    "gsc_indexation": "Status w Google Search Console API to potwierdzenie, czy Google zna i indeksuje daną stronę. Jeśli strona nie jest zaindeksowana, może być poza zasięgiem mechanizmów AI od Google.", 
    "4xx_errors": "Błędy 4xx (np. 404 Not Found) oznaczają niedziałające linki. Prowadzą one boty AI w ślepe zaułki, marnując budżet na indeksowanie i sygnalizując, że strona jest słabo utrzymana. Modele AI mogą nie odwoływać się do niepewnych, niedziałających źródeł.", 
    "3xx_redirects": "Wewnętrzne przekierowania spowalniają pracę robotów, zużywając niepotrzebnie zasoby na podążanie za łańcuchem odnośników. Preferowane są bezpośrednie linki do docelowych zasobów.", 
    "meta_description": "Meta descriptions dostarczają modelom AI zwięzłego podsumowania zawartości strony. Pomaga to w szybszym zrozumieniu kontekstu i może być wykorzystane do generowania fragmentów odpowiedzi.",
    "js_content": "Roboty wyszukiwarek i AI preferują treści dostępne bezpośrednio w kodzie HTML. Jeśli kluczowe informacje pojawiają się dopiero po wykonaniu skryptów JavaScript, może to opóźnić ich indeksowanie lub, w przypadku mniej zaawansowanych botów, całkowicie uniemożliwić ich odczytanie.", 
    "Czy strona jest dodana w Google Search Console?": "GSC pozwala monitorować indeksację, widoczność w Google i pomaga wykrywać błędy techniczne na stronie. To nie jest czynnik rankingowy lecz ułatwia analizę widocznosci i wykrywanie błędów.",
    "Czy strona jest dodana w Bing Webmaster Tools?": "Bing Webmaster Tools daje dostęp do danych o cytowaniach w AI (AI Performance).",
    "Czy dodane jest 'Organization' schema z adresami 'sameAs' kierującymi do profili społecznościowych?": "Pomaga LLM jednoznacznie zidentyfikować encję marki i powiązać ją z zewnętrznymi źródłami, co zwiększa wiarygodność i poprawność informacji w odpowiedziach AI.",
    "Czy dodane jest 'Article' schema na wpisach blogowych?": "Ułatwia modelom zrozumienie struktury i kontekstu treści, dzięki czemu łatwiej wyciągają konkretne informacje do odpowiedzi.",
    "Czy dodane jest 'Author' schema na wpisach blogowych i podlinkowane do profili autorów?": "Wzmacnia sygnały E-E-A-T, co jest kluczowe dla LLM przy wyborze wiarygodnych źródeł, szczególnie w tematach eksperckich (np. zdrowie, finanse).",
    "Czy autorzy mają stworzone dedykowane podstrony z 'ProfilePage' schema?": "Pozwala modelom lepiej zrozumieć, kim są autorzy i jakie mają kompetencje.",
    "Czy dodane jest 'Breadcrumb' schema?": "Pomaga modelom zrozumieć kontekst i hierarchię strony.",
    "Czy w ustawieniach kanału na Youtube włączona jest opcja \"Zezwalaj firmom zewnętrznym na trenowanie modeli AI przy użyciu treści z mojego kanału\"?": "Umożliwia wykorzystanie treści wideo przez systemy AI, co zwiększa szansę na ich uwzględnienie w odpowiedziach LLM i rozszerza obecność marki poza stroną.",
}

# --- HELPER FUNCTIONS ---
def read_data_file(file):
    if file is None: return None
    filename = file.name
    content = file.getvalue()
    
    if filename.lower().endswith('.csv'):
        encodings = ['utf-8-sig', 'utf-16', 'utf-8', 'latin-1']
        separators = [';', ',', '\t']
        best_df = None
        for encoding in encodings:
            for sep in separators:
                try:
                    df = pd.read_csv(io.BytesIO(content), sep=sep, encoding=encoding)
                    if best_df is None or len(df.columns) > len(best_df.columns):
                        best_df = df
                except: continue
        return best_df
    elif filename.lower().endswith('.xlsx'):
        try:
            return pd.read_excel(io.BytesIO(content))
        except: return None
    return None

def get_base64_img(file):
    if file is None: return None
    try:
        return f"data:image/png;base64,{base64.b64encode(file.getvalue()).decode()}"
    except: return None

def generate_html_report(tech_answers, content_answers, social_answers, lb_answers, commentary_db, robots_text, df_sf, df_ahrefs, df_senuto, df_schema, df_js, client_name, analyzed_url, gsc_img=None, ga_img=None, lb_img=None):
    html = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <title>Raport AI Readiness - {client_name}</title>
    <link href="https://fonts.googleapis.com/css2?family=Manrope:wght@300;400;500;700;800&display=swap" rel="stylesheet">
    <style>
        :root {{ --primary: #003366; --secondary: #006699; --tertiary: #FF9900; --light: #F5F7FA; --primary-light: #e6f0fa; --border-color: #ccd9e8; }}
        body {{ font-family: 'Manrope', sans-serif; background-color: var(--light); color: #2d3436; padding: 40px 20px; line-height: 1.6; margin: 0; }}
        .container {{ max-width: 960px; margin: 0 auto; background: white; padding: 60px; border-radius: 16px; border-top: 8px solid var(--primary); box-shadow: 0 15px 35px rgba(0,0,0,0.08); }}
        h1, h2, h3 {{ font-family: 'Manrope', sans-serif; color: var(--primary); font-weight: 800; letter-spacing: -0.02em; margin-top: 1.5em; }}
        h1 {{ font-size: 2.8rem; border-bottom: 2px solid var(--light); padding-bottom: 20px; margin-top: 0; }}
        h2 {{ font-size: 1.8rem; border-left: 6px solid var(--tertiary); padding-left: 15px; color: var(--secondary); }}
        h3 {{ font-size: 1.4rem; }}
        .table-wrap {{ width: 100%; overflow-x: auto; margin: 30px 0; }}
        table {{ width: 100%; border-collapse: collapse; font-size: 0.85rem; border-radius: 8px; overflow: hidden; border: 1px solid var(--border-color); }}
        thead {{ background-color: var(--primary); color: white; }}
        th {{ padding: 12px 10px; text-align: left; font-weight: 700; font-size: 0.75rem; letter-spacing: 0.05em; }}
        th.center {{ text-align: center; }}
        td {{ padding: 10px; border-bottom: 1px solid var(--border-color); vertical-align: middle; word-break: break-word; }}
        td.center {{ text-align: center; }}
        tr:nth-child(even) {{ background-color: var(--primary-light); }}
        .commentary {{ background: var(--light); border-left: 6px solid var(--secondary); padding: 20px 30px; margin: 30px 0; font-style: italic; color: var(--primary); border-radius: 0 8px 8px 0; }}
        .robots-code {{ background-color: #f0f0f0; color: #2d3436; padding: 20px; border-radius: 8px; font-family: 'Courier New', monospace; font-size: 13px; white-space: pre-wrap; line-height: 1.4; }}
        .badge-ok {{ background:#27ae60; color:white; padding:2px 8px; border-radius:4px; font-size:0.75rem; font-weight:700; }}
        .badge-warn {{ background:#e67e22; color:white; padding:2px 8px; border-radius:4px; font-size:0.75rem; font-weight:700; }}
        @media print {{
            @page {{ margin: 1.5cm; size: A4; }}
            body {{ background: white; padding: 0; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
            .container {{ box-shadow: none; padding: 20px; border-top: none; }}
            button {{ display: none !important; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div style="text-align: right; margin-bottom: 20px;">
            <button onclick="document.title=''; window.print();" style="background:var(--tertiary); color:white; border:none; padding:10px 20px; font-size:16px; border-radius:5px; cursor:pointer; font-weight: bold;">🖨️ Zapisz do PDF (Wydrukuj)</button>
        </div>
        <h1>Raport AI Readiness: {analyzed_url}</h1>
        <p><strong>Klient:</strong> {client_name}</p>
"""
    def section(title, d):
        s = f"<h2>{title}</h2>"
        for q, a in d.items():
            icon = "✅" if "tak" in str(a).lower() else "❌" if "nie" in str(a).lower() else "➡️"
            if q == "Czy linki przychodzące kierują do stron 404?" and "tak" in str(a).lower(): icon = "❌"
            s += f"<h3>{str(q).replace('(podaj link)', '').strip()}</h3>"
            if str(a).startswith('http'): s += f"<p>{icon} <a href='{a}'>{a}</a></p>"
            else: s += f"<p>{icon} {a if a else 'Nie uzupełniono'}</p>"
            if q in commentary_db: s += f"<div class='commentary'>{commentary_db[q]}</div>"
        return s

    def html_table_centered(df, title, url_cols=None, cwv_kind=None):
        """Tabela HTML z wysródkowaniem i ew. kolorowaniem CWV."""
        if df is None or df.empty: return ""
        df = df.fillna('-')
        if url_cols is None: url_cols = []
        rows_html = ""
        header_html = ""
        for col in df.columns:
            is_url = any(u.lower() in str(col).lower() for u in ['url', 'address', 'adres'])
            css = "" if is_url or col in url_cols else " class='center'"
            header_html += f"<th{css.replace('class=', ' class=')}>{col}</th>"
        for _, row in df.iterrows():
            row_html = ""
            for col, val in zip(df.columns, row):
                is_url = any(u.lower() in str(col).lower() for u in ['url', 'address', 'adres'])
                css = " class='center'" if not is_url and col not in url_cols else ""
                
                disp_val = val
                if cwv_kind and not is_url:
                    try:
                        v = float(str(val).replace(',', '.'))
                        bg = None
                        if cwv_kind == 'LCP':
                            bg = '#27ae60' if v <= 2500 else '#f39c12' if v <= 4000 else '#e74c3c'
                        elif cwv_kind == 'CLS':
                            bg = '#27ae60' if v <= 0.1 else '#f39c12' if v <= 0.25 else '#e74c3c'
                        elif cwv_kind in ('INP', 'FCP'):
                            bg = '#27ae60' if v <= 200 else '#f39c12' if v <= 500 else '#e74c3c'
                        if bg: disp_val = f"<span style='background:{bg}; color:white; padding:3px 8px; border-radius:4px; font-weight:bold;'>{val}</span>"
                    except: pass
                
                row_html += f"<td{css}>{disp_val}</td>"
            rows_html += f"<tr>{row_html}</tr>"
            
        link_html = ""
        if cwv_kind == 'LCP': link_html = " <br><br><a href='https://web.dev/articles/lcp?hl=pl' style='font-size:0.75em; text-decoration:none;' target='_blank'>[📚 web.dev/lcp]</a>"
        elif cwv_kind == 'CLS': link_html = " <br><br><a href='https://web.dev/articles/cls?hl=pl' style='font-size:0.75em; text-decoration:none;' target='_blank'>[📚 web.dev/cls]</a>"
        elif cwv_kind in ('INP', 'FCP'): link_html = " <br><br><a href='https://web.dev/articles/inp?hl=pl' style='font-size:0.75em; text-decoration:none;' target='_blank'>[📚 web.dev/inp]</a>"
        
        return f"<h3>{title}</h3><div class='table-wrap'><table><thead><tr>{header_html}</tr></thead><tbody>{rows_html}</tbody></table></div>{link_html}"

    if gsc_img or ga_img:
        html += "<h2>1. Analiza widoczności i ruchu</h2>"
        if gsc_img:
            html += "<h3>1.1. Widoczność w Google Search Console</h3>"
            b64 = get_base64_img(gsc_img)
            if b64: html += f"<div style='margin:20px 0;'><img src='{b64}' style='max-width:100%; border-radius:8px; border:1px solid var(--border-color);'></div>"
        if ga_img:
            html += "<h3>1.2. Ruch z LLM w Google Analytics 4</h3>"
            b64 = get_base64_img(ga_img)
            if b64: html += f"<div style='margin:20px 0;'><img src='{b64}' style='max-width:100%; border-radius:8px; border:1px solid var(--border-color);'></div>"

    html += section("2. Crawling i Indeksowanie", tech_answers)
    if robots_text:
        robots_clean = "\n".join([line for line in robots_text.replace('\r', '').split('\n') if line.strip()])
        html += f"<h3>Zawartość pliku robots.txt:</h3><div class='robots-code'>{robots_clean}</div>"

    html += section("2. Treści", content_answers)
    html += section("3. Social Media", social_answers)
    html += section("4. Linkbuilding", lb_answers)
    if lb_img:
        html += "<h3>4.1. Profil linków (Ahrefs)</h3>"
        b64 = get_base64_img(lb_img)
        if b64: html += f"<div style='margin:20px 0;'><img src='{b64}' style='max-width:100%; border-radius:8px; border:1px solid var(--border-color);'></div>"

    html += "<h2>5. Analiza potencjału w AI Overviews</h2>"
    if df_ahrefs is not None and len(df_ahrefs.columns) > 1 and 'Current URL inside' in df_ahrefs.columns:
        df_ai = df_ahrefs[df_ahrefs['Current URL inside'].astype(str).str.contains('AI Overview', case=False, na=False)].sort_values(by='Volume', ascending=False)
        disp = df_ai[['Keyword', 'Volume', 'Current position', 'Current URL']].rename(columns={'Keyword': 'Słowo kluczowe', 'Volume': 'Wolumen', 'Current position': 'Pozycja organiczna', 'Current URL': 'URL'}).head(10)
        html += html_table_centered(disp, "Widoczność AI Overview - Ahrefs")

    if df_senuto is not None:
        scols = ['Słowo kluczowe', 'Pozycja organiczna', 'Najlepsza pozycja w AIO', 'URL najlepszej pozycji w AIO']
        avail = [c for c in scols if c in df_senuto.columns]
        if avail:
            disp = df_senuto[avail].rename(columns={'URL najlepszej pozycji w AIO': 'URL w AIO'}).head(10)
            html += html_table_centered(disp, "Widoczność AI Overview - Senuto")

    if df_sf is not None or df_schema is not None or df_js is not None:
        html += "<h2>6. Analiza techniczna (Screaming Frog)</h2>"
        if df_sf is not None:
            if 'Indexability' in df_sf.columns:
                non_idx = df_sf[df_sf['Indexability'] == 'Non-Indexable'][['Address', 'Indexability Status', 'Status Code']]
                html += html_table_centered(non_idx, f"Strony nieindeksowalne ({len(non_idx)})")
            cwv_cols = ['Largest Contentful Paint Time (ms)', 'Cumulative Layout Shift']
            # Screaming Frog może mieć FCP lub INP (Interaction to Next Paint)
            inp_col = 'Interaction to Next Paint (ms)' if 'Interaction to Next Paint (ms)' in df_sf.columns else 'First Contentful Paint Time (ms)'
            
            # Filtrowanie tylko Status Code 200 dla CWV
            df_sf_200 = df_sf[df_sf['Status Code'] == 200] if 'Status Code' in df_sf.columns else df_sf
            
            if all(c in df_sf.columns for c in cwv_cols + ([inp_col] if inp_col in df_sf.columns else [])):
                html += "<h3>Core Web Vitals</h3>"
                html += html_table_centered(df_sf_200[['Address', 'Largest Contentful Paint Time (ms)']].sort_values(by='Largest Contentful Paint Time (ms)', ascending=False).head(5), "Najwolniejsze strony (LCP)", cwv_kind='LCP')
                html += html_table_centered(df_sf_200[['Address', 'Cumulative Layout Shift']].sort_values(by='Cumulative Layout Shift', ascending=False).head(5), "Strony z najwyższym przesunięciem (CLS)", cwv_kind='CLS')
                if inp_col in df_sf.columns:
                    label = "Interaktywność (INP)" if "Interaction" in inp_col else "Pierwsze wyrenderowanie (FCP)"
                    html += html_table_centered(df_sf_200[['Address', inp_col]].sort_values(by=inp_col, ascending=False).head(5), f"{label}", cwv_kind='INP')
            if 'Status Code' in df_sf.columns:
                err4xx = df_sf[(df_sf['Status Code'] >= 400) & (df_sf['Status Code'] < 500)]
                if not err4xx.empty: html += html_table_centered(err4xx[['Address', 'Status Code']].head(10), f"Strony zwracające błąd 4xx ({len(err4xx)})")
                err3xx = df_sf[(df_sf['Status Code'] >= 300) & (df_sf['Status Code'] < 400)]
                if not err3xx.empty:
                    cols3 = [c for c in ['Address', 'Status Code', 'Redirect URL'] if c in err3xx.columns]
                    html += html_table_centered(err3xx[cols3].head(10), f"Strony z przekierowaniem 3xx ({len(err3xx)})")
            # Meta Description
            if all(c in df_sf.columns for c in ['Status Code', 'Meta Description 1']):
                df200 = df_sf[df_sf['Status Code'] == 200][['Address', 'Meta Description 1']].copy()
                empty_md = df200[df200['Meta Description 1'].isna() | (df200['Meta Description 1'].astype(str).str.strip() == '')]
                dupl_md = df200[df200.duplicated(subset=['Meta Description 1'], keep=False) & df200['Meta Description 1'].notna() & (df200['Meta Description 1'].astype(str).str.strip() != '')]
                html += "<h3>Meta Opisy</h3>"
                if empty_md.empty:
                    html += "<p><span class='badge-ok'>✅ OK</span> Brak pustych meta description.</p>"
                else:
                    html += html_table_centered(empty_md.rename(columns={'Address': 'URL', 'Meta Description 1': 'Meta Description'}).head(20), f"Puste meta description ({len(empty_md)} stron)")
                if dupl_md.empty:
                    html += "<p><span class='badge-ok'>✅ OK</span> Brak zduplikowanych meta description.</p>"
                else:
                    html += html_table_centered(dupl_md.rename(columns={'Address': 'URL', 'Meta Description 1': 'Meta Description'}).sort_values('Meta Description').head(20), f"Zduplikowane meta description ({len(dupl_md)} stron)")

        if df_js is not None:
            js_cols_h = [c for c in ['Address', 'HTML Word Count', 'Rendered HTML Word Count', 'Word Count Change', 'JS Word Count %'] if c in df_js.columns]
            if js_cols_h:
                df_js_disp = df_js[js_cols_h].copy()
                if 'JS Word Count %' in df_js_disp.columns:
                    df_js_disp['JS Word Count %'] = df_js_disp['JS Word Count %'].round(0).astype(int, errors='ignore')
                    df_js_disp = df_js_disp.sort_values(by='JS Word Count %', ascending=False)
                html += html_table_centered(df_js_disp.head(10), "Zależność od JavaScript (Top 10)")

        if df_schema is not None and 'Indexability' in df_schema.columns and 'Address' in df_schema.columns:
            df_s = df_schema[df_schema['Indexability'] == 'Indexable'].sort_values('Address', ascending=True)
            t_cols = [c for c in df_s.columns if c.startswith('Type-')][:5]
            html += html_table_centered(df_s[['Address'] + t_cols].fillna('-').head(10), "Dane Strukturalne (Schema)")
            html += "<p style='font-size:0.9em; color:#555;'>Pełne błędy, ostrzeżenia i wszystkie wykryte typy schema dla wszystkich adresów znajdują się w dołączonym arkuszu XLSX.</p>"

    html += "<p style='margin-top:40px; font-style:italic; color:#666; border-top:1px solid var(--border-color); padding-top:20px;'>Pełne dane dotyczące błędów technicznych znajdują się w pliku XLSX.</p>"
    html += "</div></body></html>"
    return html

def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    r.font.underline = True
    return hyperlink

def add_strategic_commentary(document, key, commentary_db):
    if key in commentary_db:
        document.add_paragraph(commentary_db[key])
    document.add_paragraph()

# Main UI
st.title("🚀 AI Readiness Report Generator")
st.markdown("Automatyczne narzędzie do audytu gotowości witryny na zmiany w ekosystemie AI (SGE/AIO).")

tabs = st.tabs(["📌 Instrukcja", "📁 Dane Podstawowe", "🔧 Audyt Techniczny", "✍️ Treści i Social", "🔗 Linkbuilding", "🧠 Analiza Semantyczna", "📄 Generuj Raport"])

# --- TAB 0: Instrukcja ---
with tabs[0]:
    st.subheader("📌 Instrukcja korzystania z narzędzia")
    st.markdown("""
Wszystkie pola w formularzu są opcjonalne. Jeśli czegoś nie chcesz umieszczać w raporcie AI Readiness to nie musisz uzupełniać konkretnego pola. Im więcej pól uzupełnisz tym większą wartość przekażemy dla klienta :)
    """)
    st.divider()
    st.markdown("### Punkty 1-3 możesz uzupełnić ręcznie lub skorzystać z gotowego CONFIGU do Screaming Frog")
    
    try:
        with open("Config AI Readiness.seospiderconfig", "rb") as f:
            config_bytes = f.read()
        st.download_button(
            label="📥 Kliknij aby pobrać plik konfiguracyjny",
            data=config_bytes,
            file_name="Config_AI_Readiness.seospiderconfig",
            mime="application/octet-stream"
        )
    except Exception as e:
        st.error(f"Nie udało się załadować pliku konfiguracji: {e}")

    st.info("Plik konfiguracyjny wgrasz przez Configuration -> Profiles (na samym dole) -> Load")
    st.divider()

    st.markdown("""
**1. Włącz renderowanie JavaScript**
Configuration -> Spider -> Rendering -> JavaScript

**2. Włącz pobieranie informacji o Schema**
Configuration -> Spider -> Extraction -> na dole jest pole "Structured Data" - zaznacz wszystko poza "Case-Sensitive"

**3. Włącz sprawdzanie duplikatów w treści**
Configuration -> Content -> Duplicates -> zaznacz Enable Near Duplicates -> wpisz 90 (%)

**4. Włącz API PageSpeed Insights**
Configuration -> API Access -> PageSpeed Insights -> Source oznacz jako "Remote" i wklej klucz API (jeśli nie masz swojego, odezwij się do Jarosław Muzyka)

**5. Wklej adres do Screaming Frog, rozpocznij crawl i poczekaj aż zostanie ukończony na 100% (zarówno crawl jak i pobieranie danych przez API).**

Będziesz potrzebować danych z zakładek:

**Internal -> Oznacz filtr po lewej stronie "HTML"**
    """)
    if os.path.exists("internal-html.png"):
        st.image("internal-html.png", caption="Screaming Frog: Internal -> HTML")
    
    st.markdown("**JavaScript**")
    if os.path.exists("javascript.png"):
        st.image("javascript.png", caption="Screaming Frog: JavaScript Content Analysis")

    st.markdown("**Structured Data**")
    if os.path.exists("structured-data.png"):
        st.image("structured-data.png", caption="Screaming Frog: Structured Data")

    st.divider()
    st.markdown("""
**6. Po ukończonym crawlu możesz w górnym menu nacisnąć "Crawl Analysis" żeby wykryć czy na stronie znajdują się np duplikaty treści.**

**7. Pobierz dane z Ahrefs o widoczności w AI Overview** - [link do Ahrefs](https://app.ahrefs.com/v2-site-explorer/organic-keywords?brandedMode=all&chartGranularity=monthly&chartInterval=all&chartMetric=Keywords&compareDate=dontCompare&country=allGlobal&currentDate=today&dataMode=keywords&hiddenColumns=AllIntents%7C%7CCPC%7C%7CEntities%7C%7CKD%7C%7COtherIntents%7C%7CPaidTraffic%7C%7CPositionHistory%7C%7CSF%7C%7CUserIntents&intentsAttrs=&keywordRules=&limit=100&localMode=all&mainOnly=0&mode=subdomains&multipleUrlsOnly=0&offset=0&performanceChartTopPosition=top11_20%7C%7Ctop21_50%7C%7Ctop3%7C%7Ctop4_10%7C%7Ctop51&positionChanges=&projectId=2396345&serpMatch=%5B%22All%22%5D&serpRules=%7B%22comparisonMode%22%3A%5B%22Current%22%5D%2C%22mode%22%3A%22ranked%22%2C%22features%22%3A%5B%22ai_overview%22%5D%2C%22featuresMatchType%22%3A%5B%22All%22%5D%7D&sort=OrganicTrafficInitial&sortDirection=desc&target=oralb.pl%2F&urlRules=&volume_type=average)

**8. Pobierz dane z Senuto o widoczności w AI Overview** - [link do Senuto](https://app.senuto.com/visibility-analysis/ai-overviews?domain=oralb.pl&fetch_mode=subdomain&country_id=200)

**9. Zrób screena ze statystykami z "Backlink Profile" z Ahrefs** - [link do Ahrefs Backlinks](https://app.ahrefs.com/v2-site-explorer/overview?backlinksChartMode=metrics&backlinksChartPerformanceSources=domainRating&backlinksCompetitorsSource=%22UrlRating%22&backlinksRefdomainsSource=%22RefDomainsNew%22&bestFilter=all&brandedTrafficChartMetric=organic-traffic&brandedTrafficSource=target-brand&chartGranularity=monthly&chartInterval=year&competitors=&countries=&country=all&entitiesCategory=organisations&generalChartBrandedTraffic=non-branded&generalChartMode=metrics&generalChartPerformanceSources=organicTraffic%7C%7CorganicTrafficValue&generalCompetitorsSource=%22OrganicTraffic%22&generalCountriesSource=organic-traffic&generalEntitiesChartMetric=Traffic&generalPagesByTrafficChartMode=Percentage&generalPagesByTrafficSource=Pages%7C%7CTraffic&highlightChanges=1y&intentsMainSource=informational&keywordsSource=all&mode=subdomains&organicChartBrandedTraffic=non-branded&organicChartMode=metrics&organicChartPerformanceSources=organicTraffic&organicCompetitorsSource=%22OrganicTraffic%22&organicCountriesSource=organic-traffic&organicEntitiesChartMetric=Traffic&organicPagesByTrafficChartMode=Percentage&organicPagesByTrafficSource=Pages&overviewSerpChartMode=Own&overviewSerpChartSpec=AIOverview%7C%7CAdwordsBottom%7C%7CAdwordsTop%7C%7CDiscussions%7C%7CFeaturedSnippet%7C%7CImagePack%7C%7CKnowledgeCard%7C%7CKnowledgePanel%7C%7CLocalPack%7C%7CPaidSiteLinks%7C%7CPeopleAlsoAsk%7C%7CShoppingAds%7C%7CShoppingOrganic%7C%7CSitelinks%7C%7CThumbnail%7C%7CTopStories%7C%7CTweets%7C%7CVideoPreview%7C%7CVideos&overviewSerpManyChartSpec=Own%7C%7CTotal&overview_tab=backlinks&paidSearchPaidKeywordsByTopPositionsChartMode=Percentage&paidTrafficSources=cost%7C%7Ctraffic&projectId=2396345&target=oralb.pl%2F&topLevelDomainFilter=all&topOrganicKeywordsMode=normal&topOrganicPagesMode=normal&trafficType=Organic&volume_type=monthly)

**10. Zrób screeny z Google Search Console** - można oznaczyć moment wejścia AI Overview w Polsce żeby pokazać jak to wpłynęło na kliknięcia

**11. Zrób screen z raportu śledzenia ruchu z LLM z Google Analytics** - [instrukcja](https://toponline.pl/blog/jak-sledze-ruch-z-ai-w-ga4)
    """)

# --- TAB 1: Dane Podstawowe ---
with tabs[1]:
    col1, col2 = st.columns([1, 1])
    with col1:
        analyzed_url = st.text_input("Adres analizowanej strony:", placeholder="https://example.com", key="url_input")
        client_name = st.text_input("Nazwa Klienta:", placeholder="Firma X", key="client_input")
    with col2:
        logo_file = st.file_uploader("Logo Klienta (PNG, JPG, SVG):", type=['png', 'jpg', 'jpeg', 'svg'])
        if logo_file:
            st.image(logo_file, width=150)

# --- TAB 2: Audyt Techniczny ---
with tabs[2]:
    st.subheader("Checklista Techniczna")
    tech_q = [
        "Czy strona jest dodana w Google Search Console?",
        "Czy strona jest dodana w Bing Webmaster Tools?",
        "Czy dodany jest plik robots.txt?",
        "Czy plik robots.txt pozwala agentom LLM i robotom wyszukiwarek na crawlowanie strony?",
        "Czy strona jest możliwa do crawlowania przez wyszukiwarki?",
        "Czy strona jest indeksowalna?",
        "Czy dodana jest mapa strony XML?",
        "Czy dodawany jest znacznik <lastmod> w mapie witryny XML?",
        "Czy dodane są dedykowane mapy XML ze zdjęciami i filmami?",
        "Czy mapa strony XML jest dodana do robots.txt?",
        "Czy mapa strony XML zawiera tylko adresy z kodem 200, kanoniczne, nie zawiera adresów noindex?",
        "Czy dodany jest certyfikat SSL?",
        "Czy dodane jest 'Organization' schema z adresami 'sameAs' kierującymi do profili społecznościowych?",
        "Czy dodane jest 'Article' schema na wpisach blogowych?",
        "Czy dodane jest 'Author' schema na wpisach blogowych i podlinkowane do profili autorów?",
        "Czy autorzy mają stworzone dedykowane podstrony z 'ProfilePage' schema?",
        "Czy dodane jest 'Breadcrumb' schema?"
    ]
    tech_answers = {}
    cols = st.columns(2)
    for i, q in enumerate(tech_q):
        ans = cols[i%2].selectbox(q, ["— Wybierz —", "✅ Tak", "❌ Nie", "➡️ Do wdrożenia", "➡️ Nie dotyczy", "💬 Własny komentarz"], key=f"tech_{i}", help=commentary_db.get(q))
        if ans == "💬 Własny komentarz":
            ans = cols[i%2].text_input("📝 Twój komentarz:", key=f"tech_custom_{i}")
        clean = ans.replace("✅ ", "").replace("❌ ", "").replace("➡️ ", "").replace("💬 ", "").replace("— Wybierz —", "") if isinstance(ans, str) else ans
        tech_answers[q] = clean

    st.divider()
    st.subheader("Pliki z narzędzi")
    col1, col2 = st.columns(2)
    robots_file = col1.file_uploader("Plik robots.txt (opcjonalnie):", type=['txt'])
    sf_file = col1.file_uploader("Plik z audytem Screaming Frog (Internal All):", type=['csv', 'xlsx'])
    gsc_img = col2.file_uploader("Screen z Google Search Console:", type=['png', 'jpg', 'jpeg'])
    ga_img = col2.file_uploader("Screen z Google Analytics (Ruch AI):", type=['png', 'jpg', 'jpeg'])
    
    col3, col4 = st.columns(2)
    ahrefs_file = col3.file_uploader("Ahrefs (AI Overview):", type=['csv', 'xlsx'])
    senuto_file = col3.file_uploader("Senuto (AI Overview):", type=['csv', 'xlsx'])
    js_file = col4.file_uploader("SF (JS Content Analysis):", type=['csv', 'xlsx'])
    schema_file = col4.file_uploader("SF (Structured Data):", type=['csv', 'xlsx'])

# --- TAB 3: Treści i Social ---
with tabs[3]:
    st.subheader("Audyt Treści")
    content_q = [
        "Czy Twoje najważniejsze strony zostały zaktualizowane w ciągu ostatnich 6 miesięcy?",
        "Czy widoczne są daty publikacji i ostatniej modyfikacji?",
        "Czy podstrony zawierają unikalną treść?",
        "Czy kluczowe strony zawierają sekcje FAQ?",
        "Czy dodane są linki do źródeł naukowych, raportów branżowych albo źródeł pierwotnych?"
    ]
    content_answers = {}
    for q in content_q:
        ans = st.selectbox(q, ["— Wybierz —", "✅ Tak", "❌ Nie", "➡️ Częściowo", "💬 Własny komentarz"], key=f"cont_{q}", help=commentary_db.get(q, commentary_db.get("tresci_general")))
        if ans == "💬 Własny komentarz":
            ans = st.text_input("📝 Twój komentarz:", key=f"cont_custom_{q}")
        clean = ans.replace("✅ ", "").replace("❌ ", "").replace("➡️ ", "").replace("💬 ", "").replace("— Wybierz —", "") if isinstance(ans, str) else ans
        content_answers[q] = clean

    st.divider()
    st.subheader("Social Media")
    social_q = [
        "Czy marka ma utworzony profil społecznościowy na Facebook? (podaj link)",
        "Czy marka ma utworzony profil społecznościowy na Instagram? (podaj link)",
        "Czy marka ma utworzony profil społecznościowy na Tiktok? (podaj link)",
        "Czy marka ma utworzony profil społecznościowy na Youtube? (podaj link)",
        "Czy w ciągu ostatniego miesiaca został dodany materiał na Facebook?",
        "Czy w ciągu ostatniego miesiaca został dodany materiał na Instagram?",
        "Czy w ciągu ostatniego miesiaca został dodany materiał na Tiktok?",
        "Czy w ciągu ostatniego miesiaca został dodany materiał na Youtube?",
        "Czy w ustawieniach kanału na Youtube włączona jest opcja \"Zezwalaj firmom zewnętrznym na trenowanie modeli AI przy użyciu treści z mojego kanału\"?"
    ]
    social_answers = {}
    for q in social_q:
        if "trenowanie modeli AI przy użyciu treści" in q:
            st.markdown(f"**{q}**\n👉 [Instrukcja jak to zrobić](https://support.google.com/youtube/answer/15509944?sjid=13268001827056761517-EU&hl=pl)")
            ans = st.selectbox("Odpowiedź:", ["— Wybierz —", "✅ Tak", "❌ Nie", "💬 Własny komentarz"], key=f"soc_{q}", help=commentary_db.get(q))
            if ans == "💬 Własny komentarz":
                ans = st.text_input("📝 Twój komentarz:", key=f"soc_custom_{q}")
        elif "podaj link" in q:
            social_answers[q] = st.text_input(q, key=f"soc_{q}", help=commentary_db.get(q, commentary_db.get("social_media_general")))
            continue
        else:
            ans = st.selectbox(q, ["— Wybierz —", "✅ Tak", "❌ Nie", "💬 Własny komentarz"], key=f"soc_{q}", help=commentary_db.get(q, commentary_db.get("social_media_general")))
            if ans == "💬 Własny komentarz":
                ans = st.text_input("📝 Twój komentarz:", key=f"soc_custom_{q}") # Fix key conflict if any
        
        clean = ans.replace("✅ ", "").replace("❌ ", "").replace("➡️ ", "").replace("💬 ", "").replace("— Wybierz —", "") if isinstance(ans, str) else ans
        social_answers[q] = clean

# --- TAB 4: Linkbuilding ---
with tabs[4]:
    st.subheader("Audyt Profilu Linków")
    lb_q = [
        "Czy autorytet strony wyrażony DR rośnie lub jest stabilny?",
        "Czy stronie stale przybywa linków przychodzących?",
        "Czy linki przychodzące kierują do stron 404?"
    ]
    lb_answers = {}
    for q in lb_q:
        ans = st.selectbox(q, ["— Wybierz —", "✅ Tak", "❌ Nie", "➡️ Wymaga analizy", "💬 Własny komentarz"], key=f"lb_{q}", help=commentary_db.get(q, commentary_db.get("linkbuilding_general")))
        if ans == "💬 Własny komentarz":
            ans = st.text_input("📝 Twój komentarz:", key=f"lb_custom_{q}")
        clean = ans.replace("✅ ", "").replace("❌ ", "").replace("➡️ ", "").replace("💬 ", "").replace("— Wybierz —", "") if isinstance(ans, str) else ans
        lb_answers[q] = clean
    lb_img = st.file_uploader("Screen z Ahrefs (Backlink profile):", type=['png', 'jpg', 'jpeg'])

# --- TAB 5: Analiza Semantyczna ---
with tabs[5]:
    st.subheader("🧠 Analiza Semantyczna Pojedynczego Wpisu")
    st.markdown("""Ta zakładka opisuje **dodatkową, płatną usługę audytu Content Intelligence** — pogłębioną analizę pojedynczego artykułu pod kątem algorytmów Google i modeli AI takich jak ChatGPT czy Google AI Overviews.
    
> 💡 Do indywidualnej wyceny.""", unsafe_allow_html=False)
    st.divider()

    with st.expander("🎯 Wstęp: Jak Google i AI czytają Twoje treści?", expanded=True):
        st.markdown("""
**Celem tego audytu** nie jest tylko sprawdzenie, czy tekst "dobrze się czyta" ludziom. Naszym głównym zadaniem jest dostosowanie treści do sposobu, w jaki analizują ją **algorytmy Google oraz nowoczesne modele AI** (takie jak ChatGPT czy Google AI Overviews).

Audyt składa się z **3 głównych filarów**, które sprawdzają treść pod kątem co najmniej 12 kluczowych kryteriów:

1. **Zgodność z Central Search Intent (CSI)** – czy algorytm rozumie, o czym dokładnie piszesz i dla kogo?
2. **Jakość treści** – jak kosztowna i trudna w interpretacji jest Twoja strona dla robota?
3. **Ocena E-E-A-T** – czy Google uważa Cię za wiarygodnego eksperta?
""")
        try:
            with open("example/Przykladowa-analiza-semantyczna-tresci.html", "rb") as fh:
                st.download_button("📄 Pobierz przykład gotowej analizy semantycznej", data=fh.read(), file_name="Przykladowa-analiza-semantyczna.html", mime="text/html")
        except: pass

    with st.expander("🔍 Filar 1: Zgodność z Central Search Intent (CSI)"):
        st.markdown("""
*(Analiza: EAV GAP + BLUF + Chunk + URR)*

Tutaj sprawdzamy, czy Twój artykuł odpowiada na intencję użytkownika i czy jest zbudowany tak, aby maszyna mogła bezbłędnie zidentyfikować temat przewodni.

#### Central Search Intent (CSI)
To matematyczne połączenie tematu (Encji) z kontekstem źródła. Algorytm musi wiedzieć, z jakiej perspektywy opisujesz temat.

> **PRZYKŁAD:**
> - *Encja Centralna:* Szczoteczka soniczna [Marka X].
> - *Source Context:* Blog stomatologiczny vs Sklep RTV/AGD.
> - Dla bloga: "Jak poprawnie myć zęby tym modelem? Instruktaż".
> - Dla sklepu: "Porównanie ceny, specyfikacja techniczna, zakup".

#### Entity-Attribute-Value (EAV)
Google dąży do wyekstrahowania z tekstu "suchych faktów". Model EAV pozwala maszynom porównywać dane:
- **Entity:** Główny obiekt (np. Szczoteczka soniczna).
- **Attribute:** Cecha encji (np. częstotliwość drgań).
- **Value:** Konkretna dana (np. 62 000 ruchów na minutę).

#### BLUF (Bottom Line Up Front)
Najważniejsza informacja musi znaleźć się na początku. Google i AI często skanują tylko pierwsze 50 słów sekcji szukając fragmentu do zacytowania.

#### CHUNK (Fragmentacja pod RAG)
RAG (Retrieval-Augmented Generation) to technologia, dzięki której AI "uczy się" z Twojej strony. Każda sekcja pod H2 powinna być **samodzielną, wyczerpującą odpowiedzią** na dany problem.

#### URR (Unique, Root, Rare)
Hierarchia atrybutów encji:
- **UNIQUE:** Cecha wyróżniająca obiekt od wszystkich innych.
- **ROOT:** Atrybuty określające czym obiekt jest u podstaw.
- **RARE:** Detale dla ekspertów, "nice to have".
""")

    with st.expander("📊 Filar 2: Jakość treści"):
        st.markdown("""
*(Analiza: CoR + Information Density + SRL + TF-IDF)*

#### CoR (Cost of Retrieval – Koszt Wydobycia)
To wydatek obliczeniowy, jaki wyszukiwarka ponosi na przeczytanie Twojej strony.
**Zasada:** Google wybierze konkurencję, która dostarczy tę samą wiedzę "taniej" (prostszy kod, szybszy serwer, zwięzła struktura zdań).

#### Information Density (Gęstość Informacji)
Stosunek konkretnych faktów do "puchu" (ang. fluff):
- ❌ *Niska gęstość:* "Szczotkowanie zębów jest bardzo ważną czynnością, którą każdy z nas powinien wykonywać codziennie..." (Dużo słów, zero wiedzy).
- ✅ *Wysoka gęstość:* "Szczotkowanie zębów 2× dziennie przez 2 minuty usuwa płytkę nazębną i zapobiega próchnicy." (Liczby, skutki, procesy).

#### SRL (Semantic Role Labeling)
Gramatyka dla robotów. Wskazuje: **Kto? Co robi? Komu?**
- ❌ *Niejasne:* "Zostało potwierdzone, że usuwanie osadu jest skuteczniejsze."
- ✅ *SRL Optimized:* "Szczoteczka soniczna usuwa płytkę nazębną o 74% skuteczniej niż szczoteczka manualna."

#### TF-IDF (Trafność terminologiczna)
Wskaźnik oceniający unikalność słownictwa:
- **TF (Term Frequency):** Ile razy słowo występuje u Ciebie.
- **IDF (Inverse Document Frequency):** Jak rzadkie/specjalistyczne jest to słowo w internecie.

Fachowa terminologia ("strefy retencji", "abrazja") daje silny sygnał ekspertyzy.
""")

    with st.expander("⭐ Filar 3: Ocena E-E-A-T"):
        st.markdown("""
*(Experience, Expertise, Authoritativeness, Trustworthiness)*

System, którym Google ocenia wiarygodność Twoją i Twojej strony. W branżach YMYL (Twoje Pieniądze lub Twoje Życie) jest to kryterium krytyczne.

| Wymiar | Co sprawdzamy? |
|--------|----------------|
| **Experience** | Czy autor faktycznie *używał* opisywanego produktu? Zdjęcia własne, opis odczuć, unikalne spostrzeżenia — vs przepisana specyfikacja producenta. |
| **Expertise** | Czy autor ma wiedzę formalną? Czy artykuł o leczeniu napisał specjalista czy anonimowy copywriter? |
| **Authoritativeness** | Czy inni eksperci cytują tę stronę? Czy jest ona liderem opinii w branży? |
| **Trust** | Czy strona jest bezpieczna? Dane zgodne z prawdą? Intencja: pomoc użytkownikowi czy agresywna sprzedaż? |
""")

# --- TAB 6: Generuj Raport ---
with tabs[6]:
    if st.button("🚀 GENERUJ RAPORT (DOCX + XLSX)", type="primary"):
        if not analyzed_url:
            st.error("Podaj adres URL strony!")
        else:
            with st.spinner("Przetwarzanie danych i generowanie plików..."):
                try:
                    class MockFile:
                        def __init__(self, path):
                            import os
                            self.name = os.path.basename(path)
                            with open(path, "rb") as f: self.data = f.read()
                        def getvalue(self): return self.data
                    
                    if st.session_state.get("use_example_files"):
                        try:
                            if sf_file is None: sf_file = MockFile("example/internal_html oralb.xlsx")
                            if ahrefs_file is None: ahrefs_file = MockFile("example/oralb.pl-organic-keywords-subdomains-pl--a_2026-02-16_21-13-42.csv")
                            if senuto_file is None: senuto_file = MockFile("example/analiza_widoczno_ci_raport_s_owa_kluczowe_ai_overviews___domain___2026_02_16_21_14.xlsx")
                            if schema_file is None: schema_file = MockFile("example/structured_data_all - oralb.xlsx")
                            if js_file is None: js_file = MockFile("example/javascript_all - oralb.xlsx")
                            if logo_file is None: logo_file = MockFile("example/logo-oralb.png")
                            if robots_file is None: robots_file = MockFile("example/robots.txt")
                            if lb_img is None: lb_img = MockFile("example/ahrefs-oralb.png")
                        except Exception as mock_err:
                            pass

                    # 1. Image preprocessing
                    logo_bytes = None
                    if logo_file:
                        if logo_file.name.lower().endswith('.svg'):
                            logo_bytes = cairosvg.svg2png(bytestring=logo_file.getvalue())
                        else:
                            logo_bytes = logo_file.getvalue()

                    # 2. DOCX Generation
                    doc = Document()
                    
                    # Apply brand colors to Word styles
                    for style_name, hex_color in [('Heading 1', 0x003366), ('Heading 2', 0x006699), ('Heading 3', 0xFF9900), ('Title', 0x003366), ('Subtitle', 0x006699)]:
                        try:
                            if style_name in doc.styles:
                                doc.styles[style_name].font.color.rgb = RGBColor((hex_color >> 16) & 255, (hex_color >> 8) & 255, hex_color & 255)
                        except: pass
                    # Header with Logo
                    p = doc.add_paragraph()
                    if logo_bytes:
                        try:
                            logo_io = io.BytesIO(logo_bytes)
                            p.add_run().add_picture(logo_io, width=Inches(1.5))
                        except Exception as logo_err:
                            st.warning(f"Problem z logiem: {logo_err}")
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    doc.add_heading('Raport AI Readiness', level=0)
                    doc.add_paragraph(f"Raport dla domeny: {analyzed_url}", 'Subtitle')
                    
                    doc.add_heading('Wstęp', level=1)
                    intro_text = ("Niniejszy dokument stanowi analizę gotowości witryny na zmiany w sposobie funkcjonowania wyszukiwarek internetowych, "
                                  "związane z wprowadzeniem generatywnych modeli językowych (LLM) i odpowiedzi AI (AI Overviews).")
                    doc.add_paragraph(intro_text)

                    # 1. Analiza widoczności
                    if gsc_img or ga_img:
                        doc.add_heading('1. Analiza widoczności i ruchu', level=1)
                        if gsc_img:
                            doc.add_heading('1.1. Widoczność w Google Search Console', level=2)
                            doc.add_picture(io.BytesIO(gsc_img.getvalue()), width=Inches(6.0))
                        if ga_img:
                            doc.add_heading('1.2. Ruch z LLM w Google Analytics 4', level=2)
                            doc.add_picture(io.BytesIO(ga_img.getvalue()), width=Inches(6.0))

                    # Helper for sections (Nested definition for scoped access or move outside)
                    def build_q_and_a_section(document, title, answers, commentary_db, robots_content=""):
                        document.add_heading(title, level=1)
                        for question, status in answers.items():
                            short_title = str(question).replace("(podaj link)", "").strip()
                            document.add_heading(short_title, level=2)
                            status_text = str(status).lower()
                            icon = "✅" if "tak" in status_text else "❌" if "nie" in status_text or not status_text else "➡️"
                            # Pytanie o 404 - gdy TAK to ŹLE, więc ❌
                            if question == "Czy linki przychodzące kierują do stron 404?" and "tak" in status_text:
                                icon = "❌"
                            p = document.add_paragraph()
                            if str(status).startswith('http'):
                                p.add_run(f'{icon} '); add_hyperlink(p, status, status)
                            else:
                                p.add_run(f'{icon} {status if status else "Nie podano"}')
                            
                            if question == "Czy dodany jest plik robots.txt?":
                                add_strategic_commentary(document, question, commentary_db)
                                if robots_content:
                                    p_robots = document.add_paragraph(); p_robots.add_run("Zawartość pliku robots.txt:").bold = True
                                    table = document.add_table(rows=1, cols=1)
                                    set_cell_shading(table.cell(0,0), "F0F0F0")
                                    p = table.cell(0,0).paragraphs[0]
                                    p.paragraph_format.space_after = Pt(0)
                                    p.paragraph_format.line_spacing = 1.0
                                    robots_clean = "\n".join([line for line in robots_content.replace('\r', '').split('\n') if line.strip()])
                                    run = p.add_run(robots_clean)
                                    run.font.name = 'Courier New'
                                    run.font.size = Pt(9)
                            else:
                                add_strategic_commentary(document, question, commentary_db)

                    def add_styled_table(document, df, title, cwv_kind=None):
                        document.add_heading(title, level=3)
                        if df is None or df.empty:
                            document.add_paragraph("Nie znaleziono danych spełniających kryteria."); return
                        df = df.fillna('-')
                        table = document.add_table(rows=1, cols=len(df.columns)); table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for i, column_name in enumerate(df.columns):
                            p = hdr_cells[i].paragraphs[0]; run = p.add_run(str(column_name)); run.font.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); set_cell_shading(hdr_cells[i], "003366")
                        for index, row in df.head(10).iterrows(): # Limit to top 10 for visibility
                            row_cells = table.add_row().cells
                            for i, (col_name, cell_value) in enumerate(zip(df.columns, row)):
                                val_str = str(cell_value)
                                bg = None
                                if cwv_kind and 'Address' not in str(col_name) and 'URL' not in str(col_name):
                                    try:
                                        v = float(val_str.replace(',', '.'))
                                        if cwv_kind == 'LCP': bg = '27ae60' if v <= 2500 else 'f39c12' if v <= 4000 else 'e74c3c'
                                        elif cwv_kind == 'CLS': bg = '27ae60' if v <= 0.1 else 'f39c12' if v <= 0.25 else 'e74c3c'
                                        elif cwv_kind in ('INP', 'FCP'): bg = '27ae60' if v <= 200 else 'f39c12' if v <= 500 else 'e74c3c'
                                    except: pass
                                row_cells[i].text = val_str
                                if bg: set_cell_shading(row_cells[i], bg)
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in p.runs:
                                        if not run.font.bold: run.font.size = Pt(8)
                                        
                        if cwv_kind == 'LCP': document.add_paragraph(); p = document.add_paragraph(); p.add_run('🔗 https://web.dev/articles/lcp?hl=pl').font.italic = True
                        elif cwv_kind == 'CLS': document.add_paragraph(); p = document.add_paragraph(); p.add_run('🔗 https://web.dev/articles/cls?hl=pl').font.italic = True
                        elif cwv_kind in ('INP', 'FCP'): document.add_paragraph(); p = document.add_paragraph(); p.add_run('🔗 https://web.dev/articles/inp?hl=pl').font.italic = True
                        
                        document.add_paragraph()

                    # Sections
                    robots_text = robots_file.getvalue().decode('utf-8', errors='ignore') if robots_file else ""
                    build_q_and_a_section(doc, '2. Crawling i Indeksowanie', tech_answers, commentary_db, robots_text)
                    
                    doc.add_heading('3. Treści', level=1)
                    add_strategic_commentary(doc, 'tresci_general', commentary_db)
                    for q, a in content_answers.items():
                        doc.add_heading(str(q), level=2)
                        status_icon = "✅" if "tak" in str(a).lower() else "❌" if "nie" in str(a).lower() else "➡️"
                        doc.add_paragraph(f"{status_icon} {a}")
                    
                    doc.add_heading('4. Social Media', level=1)
                    add_strategic_commentary(doc, 'social_media_general', commentary_db)
                    for q, a in social_answers.items():
                        doc.add_heading(str(q).replace("(podaj link)", "").strip(), level=2)
                        p = doc.add_paragraph()
                        if str(a).startswith('http'):
                            p.add_run("🔗 "); add_hyperlink(p, str(a), str(a))
                        else:
                            status_icon = "✅" if "tak" in str(a).lower() else "➡️"
                            p.add_run(f"{status_icon} {a}")

                    # Linkbuilding Section
                    doc.add_heading('5. Linkbuilding', level=1)
                    add_strategic_commentary(doc, 'linkbuilding_general', commentary_db)
                    for q, a in lb_answers.items():
                        doc.add_heading(str(q), level=2)
                        status_icon = "✅" if "tak" in str(a).lower() else "❌" if "nie" in str(a).lower() else "➡️"
                        if q == "Czy linki przychodzące kierują do stron 404?" and "tak" in str(a).lower():
                            status_icon = "❌"
                        doc.add_paragraph(f"{status_icon} {a}")
                    
                    if lb_img:
                        doc.add_heading('5.1. Profil linków (Ahrefs)', level=2)
                        doc.add_picture(io.BytesIO(lb_img.getvalue()), width=Inches(6.0))

                    # Ahrefs / Senuto Data
                    doc.add_heading('6. Analiza potencjału w AI Overviews', level=1)
                    if ahrefs_file:
                        df_ahrefs = read_data_file(ahrefs_file)
                        if df_ahrefs is not None:
                            doc.add_heading('6.1. Widoczność AI Overview - Ahrefs', level=2)
                            if len(df_ahrefs.columns) > 1 and 'Current URL inside' in df_ahrefs.columns:
                                df_ahrefs_ai = df_ahrefs[df_ahrefs['Current URL inside'].astype(str).str.contains('AI Overview', case=False, na=False)].sort_values(by='Volume', ascending=False)
                                ahrefs_disp = df_ahrefs_ai[['Keyword', 'Volume', 'Current position', 'Current URL']].rename(columns={'Keyword': 'Słowo kluczowe', 'Volume': 'Wolumen', 'Current position': 'Pozycja organiczna', 'Current URL': 'URL'})
                                add_styled_table(doc, ahrefs_disp.head(10), "")
                    
                    if senuto_file:
                        df_senuto = read_data_file(senuto_file)
                        if df_senuto is not None:
                            doc.add_heading('6.2. Widocznosć AI Overview - Senuto', level=2)
                            senuto_cols = ['Słowo kluczowe', 'Pozycja organiczna', 'Najlepsza pozycja w AIO', 'URL najlepszej pozycji w AIO']
                            available_senuto = [c for c in senuto_cols if c in df_senuto.columns]
                            if available_senuto:
                                senuto_disp = df_senuto[available_senuto].rename(columns={'URL najlepszej pozycji w AIO': 'URL w AIO'})
                                add_styled_table(doc, senuto_disp.head(10), "")

                    # Screaming Frog Data
                    if sf_file:
                        df_sf = read_data_file(sf_file)
                        if df_sf is not None:
                            doc.add_heading('7. Analiza techniczna (Screaming Frog)', level=1)
                            
                            # non-indexable
                            if 'Indexability' in df_sf.columns:
                                non_idx = df_sf[df_sf['Indexability'] == 'Non-Indexable'][['Address', 'Indexability Status', 'Status Code']]
                                doc.add_heading('7.1. Strony nieindeksowalne', level=2)
                                add_strategic_commentary(doc, 'non_indexable', commentary_db)
                                add_styled_table(doc, non_idx, f"Strony nieindeksowalne ({len(non_idx)})")
                            
                            # CWV
                            cwv_cols = ['Largest Contentful Paint Time (ms)', 'Cumulative Layout Shift']
                            inp_col = 'Interaction to Next Paint (ms)' if 'Interaction to Next Paint (ms)' in df_sf.columns else 'First Contentful Paint Time (ms)'
                            
                            # Filtrowanie tylko Status Code 200 dla CWV
                            df_sf_200 = df_sf[df_sf['Status Code'] == 200] if 'Status Code' in df_sf.columns else df_sf
                            
                            if all(c in df_sf.columns for c in cwv_cols):
                                doc.add_heading('7.2. Szybkość strony (Core Web Vitals)', level=2)
                                add_strategic_commentary(doc, 'core_web_vitals', commentary_db)
                                
                                lcp_df = df_sf_200[['Address', 'Largest Contentful Paint Time (ms)']].sort_values(by='Largest Contentful Paint Time (ms)', ascending=False).head(5)
                                add_styled_table(doc, lcp_df, "Najwolniejsze strony (LCP)", cwv_kind='LCP')
                                
                                cls_df = df_sf_200[['Address', 'Cumulative Layout Shift']].sort_values(by='Cumulative Layout Shift', ascending=False).head(5)
                                add_styled_table(doc, cls_df, "Strony z najwyższym przesunięciem (CLS)", cwv_kind='CLS')
                                
                                if inp_col in df_sf.columns:
                                    inp_label = "Interaktywność (INP)" if "Interaction" in inp_col else "Pierwsze ładowanie treści (FCP)"
                                    inp_df = df_sf_200[['Address', inp_col]].sort_values(by=inp_col, ascending=False).head(5)
                                    add_styled_table(doc, inp_df, f"Najwolniejsze: {inp_label}", cwv_kind='INP')

                            # Errors 4xx
                            if 'Status Code' in df_sf.columns:
                                err4xx = df_sf[(df_sf['Status Code'] >= 400) & (df_sf['Status Code'] < 500)]
                                if not err4xx.empty:
                                    doc.add_heading('7.3. Błędy 4xx', level=2)
                                    add_strategic_commentary(doc, '4xx_errors', commentary_db)
                                    add_styled_table(doc, err4xx[['Address', 'Status Code']].head(10), f"Strony zwracające błąd 4xx ({len(err4xx)})")
                                
                                err3xx = df_sf[(df_sf['Status Code'] >= 300) & (df_sf['Status Code'] < 400)]
                                if not err3xx.empty:
                                    doc.add_heading('7.4. Przekierowania 3xx', level=2)
                                    add_strategic_commentary(doc, '3xx_redirects', commentary_db)
                                    add_styled_table(doc, err3xx[['Address', 'Redirect URL']].head(10), f"Strony z przekierowaniem 3xx ({len(err3xx)})")

                    # Zaleznosc od JS (DOCX)
                    if js_file:
                        df_js = read_data_file(js_file)
                        if df_js is not None:
                            js_cols_docx = [c for c in ['Address', 'HTML Word Count', 'Rendered HTML Word Count', 'Word Count Change', 'JS Word Count %'] if c in df_js.columns]
                            if js_cols_docx:
                                doc.add_heading('7.5. Zależność od JavaScript', level=2)
                                add_strategic_commentary(doc, 'js_content', commentary_db)
                                df_js_docx = df_js[js_cols_docx].copy()
                                if 'JS Word Count %' in df_js_docx.columns:
                                    df_js_docx['JS Word Count %'] = df_js_docx['JS Word Count %'].round(0).astype(int, errors='ignore')
                                    df_js_docx = df_js_docx.sort_values(by='JS Word Count %', ascending=False)
                                add_styled_table(doc, df_js_docx.head(10), "Top 10 stron z najwyższą zależnością od JS")

                    # Meta Description (DOCX)
                    if sf_file:
                        df_sf_meta = read_data_file(sf_file)
                        if df_sf_meta is not None and all(c in df_sf_meta.columns for c in ['Status Code', 'Meta Description 1']):
                            doc.add_heading('7.6. Analiza Meta Description', level=2)
                            df200m = df_sf_meta[df_sf_meta['Status Code'] == 200][['Address', 'Meta Description 1']].copy()
                            empty_md = df200m[df200m['Meta Description 1'].isna() | (df200m['Meta Description 1'].astype(str).str.strip() == '')]
                            dupl_md = df200m[df200m.duplicated(subset=['Meta Description 1'], keep=False) & df200m['Meta Description 1'].notna() & (df200m['Meta Description 1'].astype(str).str.strip() != '')]
                            if not empty_md.empty:
                                add_styled_table(doc, empty_md.rename(columns={'Address': 'URL', 'Meta Description 1': 'Meta Description'}).head(20), f"Puste Meta Description ({len(empty_md)} stron)")
                            else:
                                doc.add_paragraph('✅ Brak pustych meta description.')
                            if not dupl_md.empty:
                                add_styled_table(doc, dupl_md.rename(columns={'Address': 'URL', 'Meta Description 1': 'Meta Description'}).sort_values('Meta Description').head(20), f"Zduplikowane Meta Description ({len(dupl_md)} stron)")
                            else:
                                doc.add_paragraph('✅ Brak zduplikowanych meta description.')

                    if schema_file:
                        df_schema = read_data_file(schema_file)
                        if df_schema is not None:
                            doc.add_heading('7.5. Dane Strukturalne (Schema)', level=2)
                            add_strategic_commentary(doc, 'schema_data', commentary_db)
                            
                            if 'Indexability' in df_schema.columns and 'Address' in df_schema.columns:
                                df_s = df_schema[df_schema['Indexability'] == 'Indexable'].sort_values('Address', ascending=True)
                                type_cols = [c for c in df_s.columns if c.startswith('Type-')][:5]
                                cols_to_show = ['Address'] + type_cols
                                schema_disp = df_s[cols_to_show].fillna('-')
                                add_styled_table(doc, schema_disp.head(10), "Znalezione elementy Schema (Top 10)")

                    doc.add_paragraph("\nPełne dane dotyczące błędów technicznych znajdują się w pliku XLSX.").runs[0].font.italic = True

                    # 3. XLSX Generation
                    xlsx_buffer = io.BytesIO()
                    with pd.ExcelWriter(xlsx_buffer, engine='openpyxl') as writer:
                        # 1. Summary Sheet
                        all_summary = {**tech_answers, **content_answers, **social_answers, **lb_answers}
                        pd.DataFrame(list(all_summary.items()), columns=['Pytanie', 'Odpowiedź']).to_excel(writer, sheet_name='Checklista', index=False)
                        
                        # 2. Ahrefs — tylko wybrane kolumny
                        if ahrefs_file:
                            df_ahrefs = read_data_file(ahrefs_file)
                            if df_ahrefs is not None:
                                ahrefs_keep = [c for c in ['Keyword', 'Volume', 'Current position', 'Current URL', 'Current URL inside'] if c in df_ahrefs.columns]
                                df_ahrefs[ahrefs_keep].fillna('-').to_excel(writer, sheet_name='Ahrefs_AI_Overview_Full', index=False)
                        
                        # 3. Senuto
                        if senuto_file:
                            df_senuto_f = read_data_file(senuto_file)
                            if df_senuto_f is not None:
                                if 'Widoczność' in df_senuto_f.columns:
                                    df_senuto_f = df_senuto_f.drop(columns=['Widoczność'])
                                df_senuto_f.fillna('-').to_excel(writer, sheet_name='Senuto_AI_Overview_Full', index=False)
                                
                        # 4. Screaming Frog Data Breakdown
                        if sf_file:
                            df_sf = read_data_file(sf_file)
                            if df_sf is not None:
                                # Nieindeksowalne — 4 kolumny
                                if 'Indexability' in df_sf.columns:
                                    ni_cols = [c for c in ['Address', 'Status Code', 'Indexability', 'Canonical Link Element 1'] if c in df_sf.columns]
                                    df_sf[df_sf['Indexability'] == 'Non-Indexable'][ni_cols].fillna('-').to_excel(writer, sheet_name='Nieindeksowalne', index=False)
                                # 4xx — 2 kolumny
                                if 'Status Code' in df_sf.columns:
                                    err4_cols = [c for c in ['Address', 'Status Code'] if c in df_sf.columns]
                                    df_sf[(df_sf['Status Code'] >= 400) & (df_sf['Status Code'] < 500)][err4_cols].fillna('-').to_excel(writer, sheet_name='Bledy_4xx', index=False)
                                    # 3xx — 3 kolumny
                                    err3_cols = [c for c in ['Address', 'Status Code', 'Redirect URL'] if c in df_sf.columns]
                                    df_sf[(df_sf['Status Code'] >= 300) & (df_sf['Status Code'] < 400)][err3_cols].fillna('-').to_excel(writer, sheet_name='Przekierowania_3xx', index=False)
                                # CWV
                                cwv_cols_xlsx = ['Largest Contentful Paint Time (ms)', 'Cumulative Layout Shift']
                                inp_col_xlsx = 'Interaction to Next Paint (ms)' if 'Interaction to Next Paint (ms)' in df_sf.columns else 'First Contentful Paint Time (ms)'
                                
                                # Filtrowanie tylko Status Code 200 dla CWV
                                df_sf_200 = df_sf[df_sf['Status Code'] == 200] if 'Status Code' in df_sf.columns else df_sf
                                
                                if 'Largest Contentful Paint Time (ms)' in df_sf.columns:
                                    df_sf_200[['Address', 'Largest Contentful Paint Time (ms)']].sort_values(by='Largest Contentful Paint Time (ms)', ascending=False).fillna('-').to_excel(writer, sheet_name='CWV_LCP', index=False)
                                if 'Cumulative Layout Shift' in df_sf.columns:
                                    df_sf_200[['Address', 'Cumulative Layout Shift']].sort_values(by='Cumulative Layout Shift', ascending=False).fillna('-').to_excel(writer, sheet_name='CWV_CLS', index=False)
                                if inp_col_xlsx in df_sf.columns:
                                    sn = 'CWV_INP' if 'Interaction' in inp_col_xlsx else 'CWV_FCP'
                                    df_sf_200[['Address', inp_col_xlsx]].sort_values(by=inp_col_xlsx, ascending=False).fillna('-').to_excel(writer, sheet_name=sn, index=False)
                        
                        # 5. JS Analysis — 5 kolumn, posortowane wg JS Word Count % + zaokraglone
                        if js_file:
                            df_js = read_data_file(js_file)
                            if df_js is not None:
                                js_cols = [c for c in ['Address', 'HTML Word Count', 'Rendered HTML Word Count', 'Word Count Change', 'JS Word Count %'] if c in df_js.columns]
                                df_js_out = df_js[js_cols].copy()
                                if 'JS Word Count %' in df_js_out.columns:
                                    df_js_out['JS Word Count %'] = df_js_out['JS Word Count %'].round(0).astype(int, errors='ignore')
                                    df_js_out = df_js_out.sort_values(by='JS Word Count %', ascending=False)
                                df_js_out.fillna('-').to_excel(writer, sheet_name='Zaleznosc_od_JS', index=False)
                                
                        # 6. Schema — posortowane: Indexable pierwsze
                        if schema_file:
                            df_schema = read_data_file(schema_file)
                            if df_schema is not None:
                                if 'Indexability' in df_schema.columns:
                                    df_schema_sorted = df_schema.sort_values('Indexability', ascending=True)
                                    df_schema_sorted.to_excel(writer, sheet_name='Implementacja_Schema', index=False)
                                else:
                                    df_schema.to_excel(writer, sheet_name='Implementacja_Schema', index=False)

                        # 7. Meta Description — pusty i zduplikowane
                        if sf_file:
                            df_sf_m = read_data_file(sf_file)
                            if df_sf_m is not None and all(c in df_sf_m.columns for c in ['Status Code', 'Meta Description 1']):
                                # Wymagamy Indexability jeśli chcemy po niej filtrować
                                if 'Indexability' in df_sf_m.columns:
                                    df200_xlsx = df_sf_m[(df_sf_m['Status Code'] == 200) & (df_sf_m['Indexability'] == 'Indexable')][['Address', 'Meta Description 1']].copy()
                                else:
                                    df200_xlsx = df_sf_m[df_sf_m['Status Code'] == 200][['Address', 'Meta Description 1']].copy()
                                empty_xl = df200_xlsx[df200_xlsx['Meta Description 1'].isna() | (df200_xlsx['Meta Description 1'].astype(str).str.strip() == '')]
                                dupl_xl = df200_xlsx[df200_xlsx.duplicated(subset=['Meta Description 1'], keep=False) & df200_xlsx['Meta Description 1'].notna() & (df200_xlsx['Meta Description 1'].astype(str).str.strip() != '')]
                                if not empty_xl.empty:
                                    empty_xl.rename(columns={'Address': 'URL', 'Meta Description 1': 'Meta Description'}).to_excel(writer, sheet_name='MetaDesc_Puste', index=False)
                                if not dupl_xl.empty:
                                    dupl_xl.rename(columns={'Address': 'URL', 'Meta Description 1': 'Meta Description'}).sort_values('Meta Description').to_excel(writer, sheet_name='MetaDesc_Duplikaty', index=False)

                        # Styling XLSX z kolorami brandowymi + auto-szerokosc + wyrodkowanie
                        from openpyxl.styles import Alignment
                        header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
                        header_font = Font(color="FFFFFF", bold=True)
                        for sheetname in writer.sheets:
                            ws = writer.sheets[sheetname]
                            # Naglowki
                            for cell in ws[1]:
                                cell.fill = header_fill
                                cell.font = header_font
                                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                            # Dane + auto-szerokosc
                            col_widths = {}
                            url_like_cols = set()
                            for i, cell in enumerate(ws[1], 1):
                                col_name = str(cell.value or '')
                                if any(u in col_name.lower() for u in ['address', 'url', 'redirect']):
                                    url_like_cols.add(i)
                                if sheetname == 'MetaDesc_Duplikaty' and 'meta description' in col_name.lower():
                                    url_like_cols.add(i)
                                col_widths[i] = max(int(len(col_name) * 1.5) + 2, 12)
                            for row in ws.iter_rows(min_row=2):
                                for cell in row:
                                    col_i = cell.column
                                    val_len = len(str(cell.value or ''))
                                    # Bardziej agresywne wyliczanie szerokości (1.4x + offset)
                                    col_widths[col_i] = min(max(col_widths.get(col_i, 12), int(val_len * 1.4) + 2), 110)
                                    # Kolorowanie CWV
                                    if sheetname.startswith('CWV_') and col_i == 2:
                                        try:
                                            v = float(str(cell.value).replace(',', '.'))
                                            fill_color = None
                                            if 'LCP' in sheetname:
                                                fill_color = "27ae60" if v <= 2500 else "f39c12" if v <= 4000 else "e74c3c"
                                            elif 'CLS' in sheetname:
                                                fill_color = "27ae60" if v <= 0.1 else "f39c12" if v <= 0.25 else "e74c3c"
                                            elif 'FCP' in sheetname or 'INP' in sheetname:
                                                fill_color = "27ae60" if v <= 200 else "f39c12" if v <= 500 else "e74c3c"
                                            
                                            if fill_color:
                                                cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
                                                cell.font = Font(color="FFFFFF", bold=True)
                                        except: pass
                                        
                                    if col_i in url_like_cols:
                                        cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
                                    else:
                                        cell.alignment = Alignment(horizontal='center', vertical='center')
                            for col_i, width in col_widths.items():
                                ws.column_dimensions[get_column_letter(col_i)].width = width

                    # Final export
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    xlsx_buffer.seek(0)
                    
                    st.success("Raport wygenerowany!")
                    
                    st.session_state['ready_docx'] = doc_io.getvalue()
                    st.session_state['ready_xlsx'] = xlsx_buffer.getvalue()
                    st.session_state['ready_client'] = client_name
                    
                    # Generuj HTML
                    _df_sf = read_data_file(sf_file) if sf_file else None
                    _df_ahrefs = read_data_file(ahrefs_file) if ahrefs_file else None
                    _df_senuto = read_data_file(senuto_file) if senuto_file else None
                    _df_schema = read_data_file(schema_file) if schema_file else None
                    _df_js = read_data_file(js_file) if js_file else None
                    st.session_state['ready_html'] = generate_html_report(tech_answers, content_answers, social_answers, lb_answers, commentary_db, robots_text, _df_sf, _df_ahrefs, _df_senuto, _df_schema, _df_js, client_name, analyzed_url, gsc_img, ga_img, lb_img)
                    
                except Exception as e:
                    st.error(f"Błąd podczas generowania: {e}")
                    st.code(traceback.format_exc())

    if st.session_state.get('ready_docx') and st.session_state.get('ready_xlsx'):
        st.info("""
**ℹ️ Przed wysłaniem raportów do klienta — przeczytaj to!**

Wygenerowane pliki są gotowe w ~90%. Zalecamy ostateczne przejrzenie i drobne poprawki wizualne, ponieważ automatyczny generator może sporadycznie nieprawidłowo sformatować wybrane sekcje.

- 📝 **Plik DOCX** — sprawdź formatowanie tabel, odstępy między sekcjami oraz poprawność wstawionych zdjęć (loga, screeny).
- 🌐 **Plik HTML/PDF** — otwórz w przeglądarce i użyj `Ctrl+P` → "Zapisz jako PDF". Sprawdź czy żadna tabela nie wychodzi poza krawędź strony.
- 📊 **Plik XLSX** — zawiera pełne dane techniczne. Możesz dołączyć go jako osobny załącznik do raportu.
        """)
        col1, col2, col3 = st.columns(3)
        col1.download_button(
            label="📥 Pobierz Raport DOCX",
            data=st.session_state['ready_docx'],
            file_name=f"Raport_AI_Readiness_{st.session_state.get('ready_client', 'Klient')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        col2.download_button(
            label="📥 Pobierz Dane XLSX",
            data=st.session_state['ready_xlsx'],
            file_name=f"Dane_AI_Readiness_{st.session_state.get('ready_client', 'Klient')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        if st.session_state.get('ready_html'):
            col3.download_button(
                label="📥 Pobierz PDF/HTML",
                data=st.session_state['ready_html'],
                file_name=f"Raport_AI_Readiness_{st.session_state.get('ready_client', 'Klient')}.html",
                mime="text/html"
            )

st.sidebar.markdown(f"""
### Status Projektu
**Domena:** `{analyzed_url if analyzed_url else 'Brak'}`
**Klient:** `{client_name if client_name else 'Brak'}`
""")
if not analyzed_url:
    st.sidebar.warning("Uzupełnij dane w pierwszej zakładce!")
else:
    st.sidebar.success("Projekt zainicjowany.")

st.sidebar.divider()
st.sidebar.subheader("💡 Przykładowe dane")
st.sidebar.markdown("Zobacz jak wygląda gotowy raport lub przetestuj narzędzie używając przykładowych danych.")

def load_example_data():
    st.session_state["url_input"] = "https://oralb.pl"
    st.session_state["client_input"] = "Oral-B"
    st.session_state["use_example_files"] = True
    # Tech
    for i in range(17):
        st.session_state[f"tech_{i}"] = "✅ Tak"
    # Content
    content_questions = [
        "Czy Twoje najważniejsze strony zostały zaktualizowane w ciągu ostatnich 6 miesięcy?",
        "Czy widoczne są daty publikacji i ostatniej modyfikacji?",
        "Czy podstrony zawierają unikalną treść?",
        "Czy kluczowe strony zawierają sekcje FAQ?",
        "Czy dodane są linki do źródeł naukowych, raportów branżowych albo źródeł pierwotnych?"
    ]
    for q in content_questions:
        st.session_state[f"cont_{q}"] = "✅ Tak"
    # Social Medias
    social_questions = [
        "Czy marka ma utworzony profil społecznościowy na Facebook? (podaj link)",
        "Czy marka ma utworzony profil społecznościowy na Instagram? (podaj link)",
        "Czy marka ma utworzony profil społecznościowy na Tiktok? (podaj link)",
        "Czy marka ma utworzony profil społecznościowy na Youtube? (podaj link)",
        "Czy w ciągu ostatniego miesiaca został dodany materiał na Facebook?",
        "Czy w ciągu ostatniego miesiaca został dodany materiał na Instagram?",
        "Czy w ciągu ostatniego miesiaca został dodany materiał na Tiktok?",
        "Czy w ciągu ostatniego miesiaca został dodany materiał na Youtube?",
        "Czy w ustawieniach kanału na Youtube włączona jest opcja \"Zezwalaj firmom zewnętrznym na trenowanie modeli AI przy użyciu treści z mojego kanału\"?"
    ]
    for q in social_questions:
        if "podaj link" in q:
            st.session_state[f"soc_{q}"] = "https://facebook.com/oralb" if "Facebook" in q else "https://instagram.com/oralb"
        else:
            st.session_state[f"soc_{q}"] = "✅ Tak"
    # Linkbuilding
    lb_questions = [
        "Czy autorytet strony wyrażony DR rośnie lub jest stabilny?",
        "Czy stronie stale przybywa linków przychodzących?",
        "Czy linki przychodzące kierują do stron 404?"
    ]
    for q in lb_questions:
        st.session_state[f"lb_{q}"] = "✅ Tak"
if st.sidebar.button("✨ Uzupełnij przykładowe dane w formularzu", on_click=load_example_data):
    st.sidebar.success("Formularz wypełniony przykładowymi danymi")

try:
    with open("example/Raport_AI_Readiness_Ekspercki (19).docx", "rb") as f:
        docx_bytes = f.read()
    st.sidebar.download_button(
        label="📄 Pobierz przykładowy raport (DOCX)",
        data=docx_bytes,
        file_name="Przykladowy_Raport_AI_Readiness_OralB.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
except:
    pass

try:
    with open("example/Analiza_Techniczna_Wyniki (18).xlsx", "rb") as f:
        xlsx_bytes = f.read()
    st.sidebar.download_button(
        label="📊 Pobierz przykładowy raport (XLSX)",
        data=xlsx_bytes,
        file_name="Przykladowa_Analiza_Techniczna_OralB.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
except:
    pass